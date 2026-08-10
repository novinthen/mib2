"""Build the Cabinet-facing DOCX from canonical Markdown sources.

The generated DOCX combines MIB_2.0_EXECUTIVE_PROPOSAL.md and
TECHNICAL_ANNEXES.md. Pandoc performs semantic Markdown conversion; this
script then applies the narrative_proposal preset, explicit table geometry,
running furniture and deterministic metadata with python-docx.
"""

from __future__ import annotations

import os
import subprocess
import tempfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
PROPOSAL = HERE / "MIB_2.0_EXECUTIVE_PROPOSAL.md"
ANNEXES = HERE / "TECHNICAL_ANNEXES.md"
OUTPUT = HERE / "MIB_2.0_CABINET_SUBMISSION.docx"

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 108)
TABLE_FILL = "F4F6F9"
TABLE_BORDER = "B8C2CC"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_START_DXA = 120


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), TABLE_BORDER)


def column_widths(table):
    columns = len(table.columns)
    if columns == 1:
        return [CONTENT_DXA]
    scores = []
    for col in range(columns):
        lengths = [max(4, len(row.cells[col].text.strip())) for row in table.rows]
        scores.append(min(max(lengths), 55))
    minimum = 620 if columns >= 6 else 850
    base = minimum * columns
    if base >= CONTENT_DXA:
        minimum = CONTENT_DXA // columns
        base = minimum * columns
    remaining = CONTENT_DXA - base
    total_score = sum(scores) or columns
    widths = [minimum + int(remaining * score / total_score) for score in scores]
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def apply_table_geometry(table):
    widths = column_widths(table)
    wide_table = len(widths) >= 9
    table_font_size = 7.0 if wide_table else (8.0 if len(widths) >= 6 else 8.5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for col_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_index])
            set_cell_margins(cell, top=60 if wide_table else 80,
                             start=60 if wide_table else 120,
                             bottom=60 if wide_table else 80,
                             end=60 if wide_table else 120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), TABLE_FILL)
                cell._tc.get_or_add_tcPr().append(shading)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, size=table_font_size, color=INK, bold=(row_index == 0 or run.bold))


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def insert_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    new_para._element = new_p
    if text:
        new_para.add_run(text)
    return new_para


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
        "Heading 4": (11, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for list_name in ("List Bullet", "List Number"):
        if list_name in doc.styles:
            style = doc.styles[list_name]
            style.font.name = "Calibri"
            style.font.size = Pt(11)
            style.paragraph_format.left_indent = Inches(0.375)
            style.paragraph_format.first_line_indent = Inches(-0.194)
            style.paragraph_format.space_after = Pt(4)
            style.paragraph_format.line_spacing = 1.208

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.text = "MIB 2.0  |  Preliminary Cabinet Consideration"
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED, bold=True)
    add_page_field(section.footer.paragraphs[0])

    for table in doc.tables:
        apply_table_geometry(table)

    paragraphs = doc.paragraphs
    if paragraphs:
        title = paragraphs[0]
        title.style = doc.styles["Heading 1"]
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(90)
        title.paragraph_format.space_after = Pt(10)
        for run in title.runs:
            set_run_font(run, size=26, color=INK, bold=True)

        subtitle = insert_after(title, "Evidence-led six-year policy architecture for preliminary decision")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(16)
        for run in subtitle.runs:
            set_run_font(run, size=14, color=DARK_BLUE)

        status = insert_after(subtitle, "CONTROLLED DRAFT  |  2026 nominal ringgit  |  Stage 9 assurance release")
        status.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status.paragraph_format.space_after = Pt(12)
        for run in status.runs:
            set_run_font(run, size=10, color=MUTED, bold=True)

        boundary = insert_after(status, "This document seeks bounded preliminary decisions. It is not a legal clearance, Treasury-approved envelope, appropriation, programme launch authority, or beneficiary entitlement.")
        boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
        boundary.paragraph_format.space_before = Pt(10)
        boundary.paragraph_format.space_after = Pt(140)
        for run in boundary.runs:
            set_run_font(run, size=11, color=INK, italic=True)
        boundary.add_run().add_break(WD_BREAK.PAGE)

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "MIB 2.0 Technical Annexes":
            paragraph.paragraph_format.page_break_before = True
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.text.startswith("**"):
            paragraph.paragraph_format.keep_together = True

    props = doc.core_properties
    props.title = "MIB 2.0 Cabinet Submission"
    props.subject = "Evidence-led six-year policy architecture for preliminary Cabinet consideration"
    props.author = "MIB 2.0 Policy Project"
    props.keywords = "MIB 2.0, Cabinet, Malaysia, assurance, policy"
    fixed = datetime(2026, 8, 10, tzinfo=timezone.utc)
    props.created = fixed
    props.modified = fixed


def main():
    for source in (PROPOSAL, ANNEXES):
        if not source.exists():
            raise SystemExit(f"Missing canonical source: {source}")
    with tempfile.TemporaryDirectory(prefix="mib2-docx-") as temp_dir:
        raw_docx = Path(temp_dir) / "raw.docx"
        subprocess.run([
            "pandoc", str(PROPOSAL), str(ANNEXES), "--from=gfm", "--to=docx",
            "--standalone", "--output", str(raw_docx),
        ], check=True)
        doc = Document(raw_docx)
        style_document(doc)
        doc.save(OUTPUT)
    print(f"Built {OUTPUT}")


if __name__ == "__main__":
    main()
